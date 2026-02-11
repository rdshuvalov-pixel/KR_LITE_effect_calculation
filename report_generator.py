import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import io
import datetime

class WordReportGenerator:
    def __init__(self, calc_instance, product_id, params, results_summary=None, activation_params=None):
        self.calc = calc_instance
        self.pid = product_id
        self.params = params
        self.summary = results_summary
        self.activation_params = activation_params
        self.doc = Document()
        
        # Basic styling setup
        self._setup_styles()

    def _setup_styles(self):
        # Could customize styles here if needed
        pass

    def _add_heading(self, text, level=1):
        self.doc.add_heading(text, level)

    def _add_paragraph(self, text, bold=False):
        p = self.doc.add_paragraph()
        runner = p.add_run(text)
        if bold:
            runner.bold = True
        return p

    def _create_table_from_df(self, df, header_bg_color="E7E6E6", column_map=None):
        if df.empty:
            self.doc.add_paragraph("Нет данных.")
            return

        # Filter and rename columns if map provided
        if column_map:
            cols_to_keep = [c for c in column_map.keys() if c in df.columns]
            df = df[cols_to_keep].rename(columns=column_map)
        
        table = self.doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        
        # Header
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(df.columns):
            hdr_cells[i].text = str(col_name)
            # Add shading
            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), header_bg_color))
            hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
            # Bold header
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Rows
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                # Basic formatting
                if isinstance(val, (int, float)):
                    if abs(val) < 1 and val != 0:
                        row_cells[i].text = f"{val:.4f}" # Uplift
                    elif isinstance(val, float):
                        row_cells[i].text = f"{val:,.2f}"
                    else:
                        row_cells[i].text = str(val)
                elif isinstance(val, (pd.Timestamp, datetime.date)):
                    row_cells[i].text = val.strftime('%d.%m.%Y')
                else:
                    row_cells[i].text = str(val) if val is not None else "-"

    def generate(self):
        # 0. Header
        prod_name = self.calc.product_names.get(self.pid, f"ID {self.pid}")
        self.doc.add_heading(f"Отчет по товару: {prod_name}", 0)
        self.doc.add_paragraph(f"Дата генерации: {datetime.datetime.now().strftime('%d.%m.%Y %H:%M')}")
        self.doc.add_paragraph(f"ID товара: {self.pid}")

        # 1. Price Changes
        self._add_price_changes_section()
        
        # 2. Raw Sales
        self._add_raw_sales_section()
        
        # 3. Weekly Data & Details
        self._add_weekly_data_section()
        
        # 4. Pre-Test Selection
        self._add_pre_test_selection_section()
        
        # 5. Final Selection
        self._add_final_selection_section()
        
        # 6. Effect Calculation
        self._add_effect_calculation_section()
        
        # 8. Pilot Summary (if available)
        if self.summary:
            self._add_pilot_summary_section()

        # Save to buffer
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _add_price_changes_section(self):
        self._add_heading("1. Вводные данные: Переоценки", 1)
        hero_prices = self.calc.test_prices[self.calc.test_prices['product_id'] == self.pid].copy()
        
        if not hero_prices.empty:
            hero_prices = hero_prices.sort_values('New_Price_Start')
            
            # Metrics similar to app.py
            changes_count = len(hero_prices)
            self.doc.add_paragraph(f"Количество переоценок: {changes_count}")
            
            # Calculate metrics
            hero_prices['Change_Pct'] = 0.0
            mask_valid = hero_prices['Current_Price'] > 0
            hero_prices.loc[mask_valid, 'Change_Pct'] = (
                (hero_prices.loc[mask_valid, 'New_Price'] - hero_prices.loc[mask_valid, 'Current_Price']) /
                hero_prices.loc[mask_valid, 'Current_Price'] * 100
            )
            
            # Helper to format percentages in the dataframe itself for better display
            display_df = hero_prices[['New_Price_Start', 'New_Price', 'Current_Price', 'Change_Pct']].copy()
            # We'll let _create_table handle float formatting or do it here
            
            self._create_table_from_df(display_df, column_map={
                'New_Price_Start': 'Дата старта',
                'New_Price': 'Новая цена', 
                'Current_Price': 'Текущая цена',
                'Change_Pct': 'Изменение %'
            })
        else:
            self.doc.add_paragraph("Нет данных о переоценках.")

    def _add_raw_sales_section(self):
        self._add_heading("2. Вводные данные: Продажи (по дням)", 1)
        hero_sales = self.calc.sales[self.calc.sales['product_id'] == self.pid].copy()
        
        if not hero_sales.empty:
            hero_sales = hero_sales.sort_values('recorded_on')
            if 'cost_volume' not in hero_sales.columns:
                hero_sales['cost_volume'] = 0
            hero_sales['profit'] = hero_sales['revenue'] - hero_sales['cost_volume']
            
            display_sales = hero_sales[['recorded_on', 'price', 'quantity', 'revenue', 'cost_at_sale', 'profit']]
            self._create_table_from_df(display_sales, column_map={
                'recorded_on': 'Дата',
                'price': 'Цена продажи',
                'quantity': 'Кол-во (шт)',
                'revenue': 'Выручка',
                'cost_at_sale': 'Себест. ед.',
                'profit': 'Прибыль'
            })
        else:
            self.doc.add_paragraph("Продаж не найдено.")

    def _add_weekly_data_section(self):
        self._add_heading("3. Вводные данные: Понедельные показатели", 1)
        
        # Use existing logic from calculator to get weekly data with prices
        weekly_df = self.calc.get_product_weekly_report_data(self.pid, self.activation_params)
        
        if not weekly_df.empty:
            display_df = weekly_df[[
                'week_formatted', 'Plan_price', 'Fact_price', 'Cost_price', 
                'Суммарная выручка', 'Суммарная прибыль'
            ]]
            
            self._create_table_from_df(display_df, column_map={
                'week_formatted': 'Неделя',
                'Plan_price': 'План. цена',
                'Fact_price': 'Факт. цена', 
                'Cost_price': 'Себест.',
                'Суммарная выручка': 'Выручка (нед.)',
                'Суммарная прибыль': 'Прибыль (нед.)'
            })
            
            # --- Detailed Weekly Breakdown ---
            self._add_heading("Детализация расчетов по неделям", 2)
            
            test_weeks_df = weekly_df[weekly_df['is_test_period'] == True]
            if not test_weeks_df.empty:
                for _, row in test_weeks_df.iterrows():
                    week_start = row['week_start']
                    week_fmt = row['week_formatted']
                    
                    self._add_heading(f"Неделя: {week_fmt}", 3)
                    
                    details = self.calc.get_weekly_details(self.pid, week_start, self.activation_params)
                    
                    # 1. Plan Price
                    self._add_paragraph("1. Плановая цена", bold=True)
                    # Use a simpler text approach than markdown parser for now
                    # Strip markdown bold formatting ** **
                    clean_plan_text = details['plan_text'].replace('**', '').replace(':green[', '').replace(']', '')
                    self.doc.add_paragraph(clean_plan_text)
                    
                    # 2. Fact Price
                    self._add_paragraph("2. Фактическая цена", bold=True)
                    self.doc.add_paragraph("Транзакции, вошедшие в расчет цены:")
                    self._create_table_from_df(details['fact_transactions'][['Дата', 'Цена', 'Кол-во']])
                    
                    clean_fact_text = details['fact_text'].replace('**', '').replace(':green[', '').replace(']', '')
                    self.doc.add_paragraph(clean_fact_text)
                    
                    # 3. Revenue
                    self._add_paragraph("3. Выручка", bold=True)
                    full_df = details['full_transactions']
                    self._create_table_from_df(full_df[['Дата', 'Цена', 'Кол-во', 'Выручка']])
                    clean_rev_text = details['revenue_text'].replace('**', '').replace(':green[', '').replace(']', '')
                    self.doc.add_paragraph(clean_rev_text)
                    
                    # 4. Cost
                    self._add_paragraph("4. Себестоимость", bold=True)
                    if not details['cost_source_data'].empty:
                        self._create_table_from_df(details['cost_source_data'])
                    clean_cost_text = details['cost_text'].replace('**', '').replace(':green[', '').replace(']', '')
                    self.doc.add_paragraph(clean_cost_text)
                    
                    # 5. Profit
                    self._add_paragraph("5. Прибыль", bold=True)
                    self._create_table_from_df(full_df[['Дата', 'Выручка', 'Себестоимость ед.', 'Прибыль']])
                    clean_prof_text = details['profit_text'].replace('**', '').replace(':green[', '').replace(']', '')
                    self.doc.add_paragraph(clean_prof_text)
                    
                    # 6. Activation
                    self._add_paragraph("6. Проверка активации", bold=True)
                    clean_act_text = details['activation_text'].replace('**', '').replace(':green[', '').replace(':red[', '').replace(']', '')
                    self.doc.add_paragraph(clean_act_text)
                    
                    self.doc.add_paragraph("") # Spacer
            else:
                self.doc.add_paragraph("Нет тестовых недель.")

    def _add_pre_test_selection_section(self):
        self._add_heading("4. Выбор дотестового периода", 1)
        
        pre_test_params = {
            'pre_test_weeks_count': self.params.get('pre_test_weeks_count', 2),
            'pre_test_stock_threshold': self.params.get('pre_test_stock_threshold', 10),
            'contiguous_pre_test': self.params.get('contiguous_pre_test', True)
        }
        
        # Need date range from weekly report to match app display
        weekly_df = self.calc.get_product_weekly_report_data(self.pid, self.activation_params)
        report_min = weekly_df['week_start'].min() if not weekly_df.empty else None
        report_max = weekly_df['week_start'].max() if not weekly_df.empty else None
        
        pre_test_details = self.calc.get_pre_test_selection_details(
            self.pid, pre_test_params, report_min=report_min, report_max=report_max
        )
        
        if pre_test_details:
            df = pre_test_details['df']
            self._create_table_from_df(df)
            
            base_stock = pre_test_details.get('base_stock', 0)
            selected_stocks = pre_test_details.get('selected_stocks', [])
            
            p = self.doc.add_paragraph()
            p.add_run(f"Базовый остаток: {base_stock:,.2f}").bold = True
            if selected_stocks:
                formula = " + ".join([f"{x:.2f}" for x in selected_stocks])
                self.doc.add_paragraph(f"Формула: ({formula}) / {len(selected_stocks)} = {base_stock:.2f}")

    def _add_final_selection_section(self):
        self._add_heading("5. Финальный выбор тестовых периодов", 1)
        self.doc.add_paragraph("Исключаем транзитные недели и недели с низким стоком.")
        
        timeline = self.calc.get_product_timeline(self.pid)
        if not timeline.empty:
            # Need to apply same status logic as in app.py
            # Since we can't easily pass the activation map from outside without passing the whole map,
            # we might need to rely on what's in timeline or re-derive. 
            # Ideally timeline should have 'activation_status' if we calculated it.
            # But get_product_timeline doesn't add activation_status by default.
            
            # Re-calculate activation map for this product locally
            act_df = self.calc.get_activation_details(specific_pid=self.pid, **self.activation_params)
            act_map = {}
            if not act_df.empty:
                act_map = {(r['product_id'], r['week_start']): r['Status'] for _, r in act_df.iterrows()}
            
            timeline['activation_status'] = timeline['week_start'].apply(
                lambda w: act_map.get((self.pid, w), "")
            )
            not_our_mask = timeline['activation_status'].str.startswith('не та цена', na=False)
            not_our_test_mask = (timeline['period_label'] == 'Test') & not_our_mask
            timeline.loc[not_our_test_mask, 'period_label'] = 'NotOurPrice'
            
            # Add Display Columns
            status_trans = {
                'Pre-Test': 'Дотестовый (База)',
                'Test': 'Тестовый (Включен)',
                'LowStock_Test': 'Исключен (Мало стока)',
                'LowStock_Before': 'Исключен (Мало стока до)',
                'Transit': 'Исключен (Транзитная)',
                'NotOurPrice': 'Исключен (Не та цена)',
                'Other': 'Другое'
            }
            timeline['Статус'] = timeline['period_label'].map(status_trans).fillna('Другое')
            
            display_df = timeline[['week_formatted', 'avg_stock', 'Статус']]
            self._create_table_from_df(display_df, column_map={
                'week_formatted': 'Неделя',
                'avg_stock': 'Средний остаток'
            })

    def _add_effect_calculation_section(self):
        self._add_heading("6. Расчет эффекта", 1)
        
        use_week_values = self.params.get('test_use_week_values', True)
        calc_mode_desc = "по каждой неделе" if use_week_values else "по среднему"
        self.doc.add_paragraph(f"Режим расчета: {calc_mode_desc}")
        
        effect_df = self.calc.get_simple_effect_details(self.pid, use_week_values=use_week_values)
        
        if not effect_df.empty:
            # 6.1 Test Group
            self._add_heading("6.1 Тестовая группа", 2)
            cols = [
                'week_formatted', 'Fact_Revenue_Real', 'PreTest_Avg_Revenue', 'Uplift_Revenue_Pct',
                'Fact_Profit_Real', 'PreTest_Avg_Profit', 'Uplift_Profit_Pct'
            ]
            self._create_table_from_df(effect_df[cols], column_map={
                'week_formatted': 'Неделя',
                'Fact_Revenue_Real': 'Выручка (Test)',
                'PreTest_Avg_Revenue': 'База (Rev)',
                'Uplift_Revenue_Pct': 'Прирост % (Rev)',
                'Fact_Profit_Real': 'Прибыль (Test)',
                'PreTest_Avg_Profit': 'База (Prof)',
                'Uplift_Profit_Pct': 'Прирост % (Prof)'
            })
            
            # Weekly Details Test
            self._add_heading("Детализация (Тест)", 3)
            for _, row in effect_df.iterrows():
                w = row['week_formatted']
                rev_pct = row['Uplift_Revenue_Pct'] * 100
                prof_pct = row['Uplift_Profit_Pct'] * 100
                self.doc.add_paragraph(f"Неделя {w}: Выручка {rev_pct:+.2f}%, Прибыль {prof_pct:+.2f}%")
            
            # 6.2 Control Group
            self._add_heading("6.2 Контрольная группа и Чистый эффект", 2)
            cols_c = [
                'week_formatted', 'Control_Revenue_Real', 'Control_Avg_Revenue', 'Control_Uplift_Revenue_Pct',
                'Control_Profit_Real', 'Control_Avg_Profit', 'Control_Uplift_Profit_Pct'
            ]
            self._create_table_from_df(effect_df[cols_c], column_map={
                'week_formatted': 'Неделя',
                'Control_Revenue_Real': 'Выручка (КГ)',
                'Control_Avg_Revenue': 'База (КГ)',
                'Control_Uplift_Revenue_Pct': 'Прирост % (КГ)',
                'Control_Profit_Real': 'Прибыль (КГ)',
                'Control_Avg_Profit': 'База (КГ Prof)',
                'Control_Uplift_Profit_Pct': 'Прирост % (КГ Prof)'
            })
            
            # Net Effect Table
            self._add_heading("Итоговый чистый прирост", 3)
            cols_net = [
                'week_formatted', 'Uplift_Revenue_Pct', 'Control_Uplift_Revenue_Pct', 'Net_Effect_Revenue_Pct',
                'Uplift_Profit_Pct', 'Control_Uplift_Profit_Pct', 'Net_Effect_Profit_Pct'
            ]
            self._create_table_from_df(effect_df[cols_net], column_map={
                'week_formatted': 'Неделя',
                'Uplift_Revenue_Pct': 'Test %',
                'Control_Uplift_Revenue_Pct': 'Control %',
                'Net_Effect_Revenue_Pct': 'Net % (Rev)',
                'Uplift_Profit_Pct': 'Test % (Prof)',
                'Control_Uplift_Profit_Pct': 'Control % (Prof)',
                'Net_Effect_Profit_Pct': 'Net % (Prof)'
            })
            
            # 6.3 Absolute Effect
            self._add_heading("6.3 Абсолютный чистый эффект", 2)
            cols_abs = [
                'week_formatted', 'Net_Effect_Revenue_Pct', 'Fact_Revenue_Real', 'Net_Abs_Effect_Revenue',
                'Net_Effect_Profit_Pct', 'Fact_Profit_Real', 'Net_Abs_Effect_Profit'
            ]
            self._create_table_from_df(effect_df[cols_abs], column_map={
                'week_formatted': 'Неделя',
                'Net_Effect_Revenue_Pct': 'Net % (Rev)',
                'Fact_Revenue_Real': 'Факт (Rev)',
                'Net_Abs_Effect_Revenue': 'Абс. Эффект (Rev)',
                'Net_Effect_Profit_Pct': 'Net % (Prof)',
                'Fact_Profit_Real': 'Факт (Prof)',
                'Net_Abs_Effect_Profit': 'Абс. Эффект (Prof)'
            })
            
            # 7. Summary
            self._add_heading("7. Итоги", 1)
            total_rev = effect_df['Net_Abs_Effect_Revenue'].sum()
            total_prof = effect_df['Net_Abs_Effect_Profit'].sum()
            self._add_paragraph(f"Доп. выручка: {total_rev:,.2f} руб.", bold=True)
            self._add_paragraph(f"Доп. прибыль: {total_prof:,.2f} руб.", bold=True)

    def _add_pilot_summary_section(self):
        self._add_heading("8. Итого по всему пилоту", 1)
        
        s = self.summary
        
        # Simple text dump of summary metrics
        self._add_paragraph("Ключевые показатели эффективности (Выручка):", bold=True)
        self.doc.add_paragraph(f"Эффект: {s.get('total_abs_effect_revenue',0):,.0f} ({s.get('effect_revenue_pct',0):.2f}%)")
        self.doc.add_paragraph(f"Глобальный оборот: {s.get('global_revenue',0):,.0f}")
        
        self._add_paragraph("Ключевые показатели эффективности (Прибыль):", bold=True)
        self.doc.add_paragraph(f"Эффект: {s.get('total_abs_effect_profit',0):,.0f} ({s.get('effect_profit_pct',0):.2f}%)")
        
        self._add_paragraph("Статистика:", bold=True)
        self.doc.add_paragraph(f"Протестировано: {s.get('tested_count',0)}")
        self.doc.add_paragraph(f"Исключено: {s.get('excluded_products_count',0)}")
